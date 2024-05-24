from queue import Queue
from docx import Document
import assemblyai as aai
from openai import OpenAI
import openai
from prompts import MAIN_PROMPT, SYSTEM_PROMPT, DOCUMENTATION_PROMPT
import pyttsx3

transcript_queue = Queue()

aai.settings.api_key = "f47fc69f44914e5d973f9d2f719ba271"

client = OpenAI(api_key=api_key)
openai.api_key = api_key


system_prompt = SYSTEM_PROMPT
message_prompt = MAIN_PROMPT

class MedicalChatbot:

    def __init__(self):
        self.context = [{"role": "system", "content": message_prompt}]
        self.finished = False
        self.patient_info = {
            'Subjective': '',
            'Objective': '',
            'Analysis': '',
            'Plan': '',
            'Implementation': '',
            'Evaluation': ''
        }

        self.initial_questions_answers = ""
        self.transcriber = None
        self.transcription_active = False

    def handle_initial_questions(self):
        questions = [
            "Are you a new patient?",
            "What is your name?",
            "What is your approximate height?",
            "What is your approximate weight?",
            "Are you currently taking any medications?",
            "Have you had any recent surgeries?",
            "Do you have any known drug allergies?",
            "Finally, what are you in for today?"
        ]
        self.initial_questions_dict = {}
        last_answer = ""
        for question in questions:
            print(question)
            answer = input("You: ")
            self.initial_questions_dict[question] = answer
            self.context.append({"role": "user", "content": question + " " + answer})
            if question.startswith("Finally"):
                last_answer = answer
        return last_answer

    def should_stop(self, message):
        if "we'll see you in the office later today" in message.lower():
            self.finished = True

    def generate_response(self, message):
        self.context.append({'role': 'user', 'content': message})
        response = client.chat.completions.create(
            model="gpt-4",
            messages=self.context
        )
        assistant_message = response.choices[0].message.content
        self.context.append({'role': 'assistant', 'content': assistant_message})
        return assistant_message
    
    def update_patient_info(self, category, content):
        self.patient_info[category] += f"{content}\n"

    def get_full_conversation(self):
        return self.context

    def create_report(self):
        chat_history = self.get_full_conversation()[1:] 
        # Build the prompt for the report
        new_prompt = DOCUMENTATION_PROMPT
        
        chat_history_string = ""
        for message in chat_history:
            role = message['role'].capitalize()
            content = message['content'].replace('\n', ' ')  # Replace newlines with spaces
            chat_history_string += f"{role}: {content}\n"
        
        new_prompt += chat_history_string
        
        try:
            # Call the OpenAI API to generate the report
            response =  client.chat.completions.create(
                model="gpt-4",
                messages=[{"role" : "system", "content" : new_prompt}]
            )
            return response
        except Exception as e:
            print(f"An error occurred {str(e)}")
            return ""
    
    def extract_and_save_report(self, report_content):
        try:
            doc = Document()
            doc.add_heading('Patient Assessment Report', 0)

            # Add initial questions and answers
            for question, answer in self.initial_questions_dict.items():
                doc.add_paragraph(f"{question}: {answer}")

            # Add a separator before the rest of the report
            doc.add_paragraph("\n---\n")

            # Add the report content
            doc.add_paragraph(report_content)
            file_path = 'Patient_Assessment_Report.docx'  # Adjust path as necessary
            doc.save(file_path)
            return file_path
        except Exception as e:
            return f"An error occurred: {str(e)}"

def on_data(transcript: aai.RealtimeTranscript):
    if not transcript.text:
        return
    if isinstance(transcript, aai.RealtimeFinalTranscript):
        transcript_queue.put(transcript.text + '')
        print("User:", transcript.text, end="\r\n")
    else:
        print(transcript.text, end="\r")

def on_error(error: aai.RealtimeError):
    print("An error occurred:", error)

def speak(text):
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()

def handle_conversation():
    bot = MedicalChatbot()
    last_initial_answer = bot.handle_initial_questions()

    if last_initial_answer:
        response = bot.generate_response(last_initial_answer)
        print("Virtual Nurse:", response)
        speak(response)

    while not bot.finished:
        transcriber = aai.RealtimeTranscriber(
            on_data=on_data,
            on_error=on_error,
            sample_rate=44_100,
        )
        transcriber.connect()
        microphone_stream = aai.extras.MicrophoneStream()
        transcriber.stream(microphone_stream)

        transcript_result = transcript_queue.get()
        response = bot.generate_response(transcript_result)

        print("\nAI:", response, end="\r\n")
        speak(response)

        bot.should_stop(response)

    if bot.finished:
        report_content = bot.create_report().choices[0].message.content
        file_path = bot.extract_and_save_report(report_content)
        print(f"Report saved to: {file_path}")

if __name__ == "__main__":
    print("Hello, I'm your virtual nurse assistant. Let's start with some basic questions.")
    handle_conversation()
